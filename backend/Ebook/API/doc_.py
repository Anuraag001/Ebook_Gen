import docx
from docx import Document
from docx.shared import Inches,Cm
from docx.enum.section import WD_SECTION

def setCoverPage(doc, image=None,title=None,author=None):
    section=doc.sections[0]
    pageWidth=section.page_width
    pageHeight=section.page_height
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = 0
    doc.add_picture(image, width=pageWidth, height=pageHeight)
    doc.add_section(WD_SECTION.CONTINUOUS)

def create_document(Header=None,Footer=None,paragraphs=None):
    doc=Document()

    section=doc.sections[0]
    if Header:
        header=section.header
        header_para=header.paragraphs[0]
        header_para.text=Header
        doc.add_heading("firstheading",1)

    if paragraphs:
        for para in paragraphs:
            doc.add_paragraph(para)

    if Footer:
        footer=section.footer
        footer_para=footer.paragraphs[0]
        footer_para.text=Footer

    #doc.add_page_break()
    return doc

def save(doc):
    doc.save("testa.docx")

def add_chapter_names(doc,chapter_names):
    doc.add_heading("Chapter Names",0)
    for i in range(len(chapter_names)):
        doc.add_heading(f"Chapter{i+1}: {chapter_names[i]}",2)
    doc.add_page_break()

if __name__=='__main__':
    text_list = [
        "This is the text for the first page.",
        "This is the text for the second page.",
        "This is the text for the third page."
    ]

    header_text = "Header Text"
    footer_text = "Footer Text"

    create_document(header_text, footer_text,text_list)
    print("done")
